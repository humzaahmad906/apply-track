"""Turn an uploaded file into plain text for the extraction step."""

from __future__ import annotations

import re
from pathlib import Path

SUPPORTED = {".pdf", ".docx", ".txt", ".md"}

_BLANK_RUN = re.compile(r"\n{3,}")
_TRAILING_WS = re.compile(r"[ \t]+$", re.MULTILINE)


class UnsupportedFile(ValueError):
    pass


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _from_pdf(path)
    elif suffix == ".docx":
        text = _from_docx(path)
    elif suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        raise UnsupportedFile(
            f"Unsupported file type '{suffix}'. Supported: {', '.join(sorted(SUPPORTED))}"
        )
    return _normalise(text)


def _from_pdf(path: Path) -> str:
    import pymupdf

    with pymupdf.open(path) as doc:
        return "\n\n".join(page.get_text("text") for page in doc)


def _from_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]

    # Many resumes lay themselves out in tables, so those cells carry real
    # content rather than decoration.
    for table in doc.tables:
        for row in table.rows:
            seen: list[str] = []
            for cell in row.cells:
                value = cell.text.strip()
                # Merged cells repeat the same object across the row.
                if value and value not in seen:
                    seen.append(value)
            if seen:
                parts.append(" | ".join(seen))

    return "\n".join(parts)


def _normalise(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ")
    text = _TRAILING_WS.sub("", text)
    text = _BLANK_RUN.sub("\n\n", text)
    return text.strip()
